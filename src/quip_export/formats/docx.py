"""Export Quip thread HTML to DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from lxml import etree
from lxml import html as lhtml


def _add_inline(node: etree._Element, para: Paragraph) -> None:
    """Recursively add inline content (text + bold/italic runs) to a paragraph."""
    if node.text:
        para.add_run(node.text)

    for child in node:
        tag = child.tag if isinstance(child.tag, str) else ""
        if tag == "img":
            para.add_run("[image]")
        else:
            run = para.add_run(child.text_content())
            if tag in ("strong", "b"):
                run.bold = True
            elif tag in ("em", "i"):
                run.italic = True

        if child.tail:
            para.add_run(child.tail)


def _add_table(node: etree._Element, doc: DocumentType) -> None:
    rows = node.xpath(".//tr")
    if not rows:
        return
    max_cols = max(
        sum(1 for c in row.xpath("th|td") if isinstance(c.tag, str))
        for row in rows
    )
    if max_cols == 0:
        return
    table = doc.add_table(rows=len(rows), cols=max_cols)
    for r_idx, row in enumerate(rows):
        cells = row.xpath("th|td")
        for c_idx, cell in enumerate(cells):
            if c_idx < max_cols:
                table.cell(r_idx, c_idx).text = cell.text_content().strip()


def _walk(node: etree._Element, doc: DocumentType) -> None:
    tag = node.tag if isinstance(node.tag, str) else ""

    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        doc.add_heading(node.text_content().strip(), level=int(tag[1]))
    elif tag == "p":
        para = doc.add_paragraph()
        _add_inline(node, para)
    elif tag == "ul":
        for child in node:
            if isinstance(child.tag, str) and child.tag == "li":
                para = doc.add_paragraph(style="List Bullet")
                _add_inline(child, para)
    elif tag == "ol":
        for child in node:
            if isinstance(child.tag, str) and child.tag == "li":
                para = doc.add_paragraph(style="List Number")
                _add_inline(child, para)
    elif tag == "li":
        para = doc.add_paragraph(style="List Bullet")
        _add_inline(node, para)
    elif tag == "table":
        _add_table(node, doc)
    elif tag == "img":
        doc.add_paragraph("[image]")
    elif tag in ("br",):
        doc.add_paragraph()
    elif tag in ("div", "body", "html", "section", "article", "main"):
        for child in node:
            _walk(child, doc)
    else:
        for child in node:
            _walk(child, doc)


def export_docx(html: str, output_path: Path) -> None:
    doc = Document()
    root = lhtml.fromstring(html)
    _walk(root, doc)
    doc.save(output_path)
