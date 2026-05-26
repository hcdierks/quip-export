"""Unit tests — DOCX export (issue #6)."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from quip_export.formats.docx import export_docx


class TestExportDocxStructure:
    def test_creates_file_at_given_path(self, tmp_path, html_document):
        out = tmp_path / "report.docx"
        export_docx(html_document, out)
        assert out.exists()
        assert out.stat().st_size > 0

    def test_h1_heading_preserved(self, tmp_path, html_document):
        out = tmp_path / "out.docx"
        export_docx(html_document, out)
        doc = Document(out)
        heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
        assert any("Quarterly Report" in t for t in heading_texts)

    def test_h2_heading_preserved(self, tmp_path, html_document):
        out = tmp_path / "out.docx"
        export_docx(html_document, out)
        doc = Document(out)
        heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading 2")]
        assert any("Executive Summary" in t for t in heading_texts)

    def test_body_paragraph_present(self, tmp_path, html_document):
        out = tmp_path / "out.docx"
        export_docx(html_document, out)
        doc = Document(out)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Revenue grew" in all_text

    def test_unordered_list_items_present(self, tmp_path, html_document):
        out = tmp_path / "out.docx"
        export_docx(html_document, out)
        doc = Document(out)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "New customers" in all_text
        assert "Churn rate" in all_text

    def test_ordered_list_items_present(self, tmp_path, html_document):
        out = tmp_path / "out.docx"
        export_docx(html_document, out)
        doc = Document(out)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Increase marketing" in all_text

    def test_hyperlink_text_present(self, tmp_path, html_document):
        out = tmp_path / "out.docx"
        export_docx(html_document, out)
        doc = Document(out)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "CFO" in all_text

    def test_table_written(self, tmp_path, html_document):
        out = tmp_path / "out.docx"
        export_docx(html_document, out)
        doc = Document(out)
        assert len(doc.tables) >= 1

    def test_table_cell_content_correct(self, tmp_path, html_document):
        out = tmp_path / "out.docx"
        export_docx(html_document, out)
        doc = Document(out)
        table = doc.tables[0]
        cell_texts = [cell.text for row in table.rows for cell in row.cells]
        assert "EMEA" in cell_texts or any("EMEA" in t for t in cell_texts)


class TestExportDocxEdgeCases:
    def test_empty_html_produces_valid_docx(self, tmp_path, html_empty):
        out = tmp_path / "empty.docx"
        export_docx(html_empty, out)
        assert out.exists()
        doc = Document(out)
        assert doc is not None

    def test_malformed_html_does_not_raise(self, tmp_path, html_malformed):
        out = tmp_path / "malformed.docx"
        export_docx(html_malformed, out)
        assert out.exists()

    def test_image_tag_replaced_with_placeholder(self, tmp_path, html_with_image):
        out = tmp_path / "img.docx"
        export_docx(html_with_image, out)
        doc = Document(out)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "[image]" in all_text.lower() or "image" in all_text.lower()

    def test_write_failure_raises_io_error(self, tmp_path, html_document):
        read_only_dir = tmp_path / "ro"
        read_only_dir.mkdir()
        import stat
        read_only_dir.chmod(stat.S_IREAD | stat.S_IEXEC)
        out = read_only_dir / "out.docx"
        try:
            with pytest.raises((PermissionError, OSError)):
                export_docx(html_document, out)
        finally:
            read_only_dir.chmod(stat.S_IRWXU)

    def test_bold_text_present_in_runs(self, tmp_path, html_document):
        out = tmp_path / "bold.docx"
        export_docx(html_document, out)
        doc = Document(out)
        bold_runs = [
            run for para in doc.paragraphs for run in para.runs if run.bold
        ]
        assert len(bold_runs) >= 1

    def test_italic_text_present_in_runs(self, tmp_path, html_document):
        out = tmp_path / "italic.docx"
        export_docx(html_document, out)
        doc = Document(out)
        italic_runs = [
            run for para in doc.paragraphs for run in para.runs if run.italic
        ]
        assert len(italic_runs) >= 1
